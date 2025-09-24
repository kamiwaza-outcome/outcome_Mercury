import os
import logging
import re
from typing import Dict, Any, List, Tuple, Optional
from pathlib import Path
import tempfile
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from googleapiclient.http import MediaIoBaseDownload
import io
import pypdf
from docx import Document
import json
import hashlib

logger = logging.getLogger(__name__)

class GoogleDriveRFPExtractor:
    """
    Extract RFP data from Google Drive folders, excluding assessment.txt
    which is meant for human review only
    """

    def __init__(self):
        self.service = self._initialize_service()
        self.temp_dir = Path(tempfile.mkdtemp(prefix="drive_rfp_"))

        # Document categories for intelligent processing
        self.document_categories = {
            'main_rfp': ['rfp', 'solicitation', 'notice', 'announcement'],
            'info_doc': ['info', 'information', 'metadata', 'details', 'summary'],
            'technical_specs': ['technical', 'specifications', 'sow', 'statement_of_work', 'requirements'],
            'legal_terms': ['contract', 'terms', 'conditions', 'legal', 'clauses'],
            'evaluation': ['evaluation', 'criteria', 'scoring', 'factors', 'selection'],
            'submission': ['submission', 'instructions', 'format', 'guidelines'],
            'amendments': ['amendment', 'modification', 'change', 'update', 'revision'],
            'cost_info': ['cost', 'pricing', 'budget', 'template', 'schedule', 'price'],
            'attachments': ['attachment', 'appendix', 'exhibit', 'annex'],
            'excluded': ['assessment']  # Files to exclude from LLM context
        }

    def _initialize_service(self):
        """Initialize Google Drive API service"""
        try:
            service_account_file = os.getenv("GOOGLE_SERVICE_ACCOUNT_FILE")

            if not service_account_file or not os.path.exists(service_account_file):
                raise FileNotFoundError(f"Service account file not found: {service_account_file}")

            credentials = service_account.Credentials.from_service_account_file(
                service_account_file,
                scopes=['https://www.googleapis.com/auth/drive.readonly']
            )

            service = build('drive', 'v3', credentials=credentials)
            logger.info("Google Drive extractor service initialized successfully")
            return service
        except Exception as e:
            logger.error(f"Failed to initialize Google Drive service: {e}")
            raise

    def extract_folder_id(self, folder_url: str) -> str:
        """Extract folder ID from Google Drive URL"""
        # Handle various URL formats
        patterns = [
            r'/folders/([a-zA-Z0-9-_]+)',
            r'id=([a-zA-Z0-9-_]+)',
            r'/d/([a-zA-Z0-9-_]+)',
        ]

        for pattern in patterns:
            match = re.search(pattern, folder_url)
            if match:
                return match.group(1)

        # If no pattern matches, assume the input is already a folder ID
        if re.match(r'^[a-zA-Z0-9-_]+$', folder_url):
            return folder_url

        raise ValueError(f"Could not extract folder ID from URL: {folder_url}")

    async def extract_from_folder(self, folder_url: str, notice_id: str = None) -> Tuple[Dict[str, Any], Dict[str, Any]]:
        """
        Main extraction method - gets all RFP data from Drive folder

        Returns:
            Tuple of (documents, metadata)
            - documents: Dict of document_name -> content
            - metadata: Dict of extracted metadata
        """
        try:
            folder_id = self.extract_folder_id(folder_url)
            logger.info(f"Extracting RFP data from folder: {folder_id}")

            # List all files in the folder RECURSIVELY
            files = await self._list_folder_contents_recursive(folder_id)
            logger.info(f"Found {len(files)} files in folder (including subfolders)")

            # Categorize documents
            categorized = self._categorize_documents(files)

            # Download and process documents (excluding assessment.txt)
            documents = await self._download_and_process_documents(categorized, folder_id)

            # Extract structured metadata from info doc if available
            metadata = await self._extract_metadata(documents, categorized, notice_id)

            # Add extraction statistics
            metadata['extraction_stats'] = {
                'total_files': len(files),
                'processed_files': len(documents),
                'excluded_files': len([f for f in files if self._should_exclude(f['name'])]),
                'total_content_size': sum(len(str(content)) for content in documents.values()),
                'source': 'google_drive',
                'folder_id': folder_id
            }

            logger.info(f"Successfully extracted {len(documents)} documents from Drive folder")
            return documents, metadata

        except Exception as e:
            logger.error(f"Error extracting from Drive folder: {e}")
            raise

    async def _list_folder_contents(self, folder_id: str) -> List[Dict[str, Any]]:
        """List all files in a Google Drive folder"""
        try:
            files = []
            page_token = None

            while True:
                # Try with shared drive support
                response = self.service.files().list(
                    q=f"'{folder_id}' in parents and trashed = false",
                    spaces='drive',
                    fields='nextPageToken, files(id, name, mimeType, size, modifiedTime, description)',
                    pageToken=page_token,
                    pageSize=100,
                    supportsAllDrives=True,  # Support shared drives
                    includeItemsFromAllDrives=True  # Include shared drive items
                ).execute()

                files.extend(response.get('files', []))
                page_token = response.get('nextPageToken', None)

                if page_token is None:
                    break

            return files

        except HttpError as e:
            logger.error(f"Error listing folder contents: {e}")
            raise

    async def _list_folder_contents_recursive(self, folder_id: str, path_prefix: str = "") -> List[Dict[str, Any]]:
        """Recursively list all files in a Google Drive folder and its subfolders"""
        try:
            all_files = []

            # Get files in current folder
            items = await self._list_folder_contents(folder_id)

            for item in items:
                # If it's a folder, recursively get its contents
                if 'folder' in item.get('mimeType', ''):
                    subfolder_name = item['name']
                    logger.info(f"Exploring subfolder: {path_prefix}{subfolder_name}")

                    # Special handling for known structure folders
                    # Prioritize Source_Documents folder for RFP source files
                    if subfolder_name in ['Source_Documents', 'source_documents', 'Source Documents']:
                        logger.info(f"Found Source_Documents folder - prioritizing this for RFP extraction")
                        sub_files = await self._list_folder_contents_recursive(
                            item['id'],
                            f"{path_prefix}{subfolder_name}/"
                        )
                        # Add path prefix to help identify source documents
                        for sf in sub_files:
                            sf['path_prefix'] = f"{path_prefix}{subfolder_name}/"
                            sf['is_source_document'] = True
                        all_files.extend(sub_files)
                    elif subfolder_name not in ['Submission_Package', 'Review_Documents', 'Generated_Responses']:
                        # Recursively explore other folders except output folders
                        sub_files = await self._list_folder_contents_recursive(
                            item['id'],
                            f"{path_prefix}{subfolder_name}/"
                        )
                        for sf in sub_files:
                            sf['path_prefix'] = f"{path_prefix}{subfolder_name}/"
                        all_files.extend(sub_files)
                else:
                    # It's a file, add it to the list
                    item['path_prefix'] = path_prefix
                    all_files.append(item)
                    logger.debug(f"Found file: {path_prefix}{item['name']}")

            return all_files

        except Exception as e:
            logger.error(f"Error in recursive folder listing: {e}")
            return []

    def _categorize_documents(self, files: List[Dict[str, Any]]) -> Dict[str, List[Dict[str, Any]]]:
        """Categorize files based on their names and types"""
        categorized = {category: [] for category in self.document_categories.keys()}
        categorized['uncategorized'] = []

        for file in files:
            filename_lower = file['name'].lower()
            categorized_flag = False

            # Check each category
            for category, keywords in self.document_categories.items():
                if any(keyword in filename_lower for keyword in keywords):
                    categorized[category].append(file)
                    categorized_flag = True
                    break

            # If not categorized, add to uncategorized
            if not categorized_flag:
                categorized['uncategorized'].append(file)

        # Log categorization results
        for category, files in categorized.items():
            if files:
                logger.info(f"Category '{category}': {len(files)} files - {[f['name'] for f in files]}")

        return categorized

    def _should_exclude(self, filename: str) -> bool:
        """Check if a file should be excluded from LLM context"""
        filename_lower = filename.lower()

        # Exclude assessment.txt and similar files
        excluded_patterns = [
            'assessment.txt',
            'assessment.docx',
            'assessment.pdf',
            'human_assessment',
            'manual_assessment',
            'review_assessment'
        ]

        return any(pattern in filename_lower for pattern in excluded_patterns)

    async def _download_and_process_documents(self, categorized: Dict[str, List[Dict]], folder_id: str) -> Dict[str, str]:
        """Download and extract text from all relevant documents"""
        documents = {}

        for category, files in categorized.items():
            # Skip excluded category
            if category == 'excluded':
                logger.info(f"Skipping {len(files)} excluded files (assessment docs)")
                continue

            for file in files:
                try:
                    # Skip if should be excluded
                    if self._should_exclude(file['name']):
                        logger.info(f"Excluding {file['name']} from LLM context")
                        continue

                    # Download and process the file
                    content = await self._download_file(file['id'], file['name'], file['mimeType'])

                    if content:
                        # Create a meaningful document key
                        doc_key = f"{category}_{file['name']}"
                        documents[doc_key] = content
                        logger.info(f"Processed {file['name']}: {len(content)} characters")

                except Exception as e:
                    logger.error(f"Error processing file {file['name']}: {e}")

        return documents

    async def _download_file(self, file_id: str, filename: str, mime_type: str) -> Optional[str]:
        """Download a file from Google Drive and extract its text content"""
        try:
            # Handle Google Docs/Sheets/Slides exports
            if 'google-apps' in mime_type:
                return await self._export_google_doc(file_id, filename, mime_type)

            # Download binary files with shared drive support
            request = self.service.files().get_media(
                fileId=file_id,
                supportsAllDrives=True
            )
            file_content = io.BytesIO()
            downloader = MediaIoBaseDownload(file_content, request)

            done = False
            while not done:
                status, done = downloader.next_chunk()

            file_content.seek(0)

            # Extract text based on file type
            return self._extract_text_from_content(file_content.read(), filename, mime_type)

        except Exception as e:
            logger.error(f"Error downloading file {filename}: {e}")
            return None

    async def _export_google_doc(self, file_id: str, filename: str, mime_type: str) -> Optional[str]:
        """Export Google Docs/Sheets/Slides to text"""
        try:
            export_mime_type = 'text/plain'

            # Choose appropriate export format
            if 'spreadsheet' in mime_type:
                export_mime_type = 'text/csv'
            elif 'presentation' in mime_type:
                export_mime_type = 'text/plain'

            request = self.service.files().export_media(
                fileId=file_id,
                mimeType=export_mime_type
            )

            file_content = io.BytesIO()
            downloader = MediaIoBaseDownload(file_content, request)

            done = False
            while not done:
                status, done = downloader.next_chunk()

            file_content.seek(0)
            return file_content.read().decode('utf-8', errors='ignore')

        except Exception as e:
            logger.error(f"Error exporting Google Doc {filename}: {e}")
            return None

    def _extract_text_from_content(self, content: bytes, filename: str, mime_type: str) -> str:
        """Extract text from binary content based on file type"""
        try:
            filename_lower = filename.lower()

            # PDF files
            if 'pdf' in mime_type or filename_lower.endswith('.pdf'):
                return self._extract_pdf_text(content)

            # Word documents
            elif 'wordprocessing' in mime_type or filename_lower.endswith(('.docx', '.doc')):
                return self._extract_docx_text(content)

            # Text files
            elif 'text' in mime_type or filename_lower.endswith(('.txt', '.md', '.csv')):
                return content.decode('utf-8', errors='ignore')

            # JSON files
            elif 'json' in mime_type or filename_lower.endswith('.json'):
                data = json.loads(content.decode('utf-8', errors='ignore'))
                return json.dumps(data, indent=2)

            # Excel files - return filename and size for now
            elif 'spreadsheet' in mime_type or filename_lower.endswith(('.xlsx', '.xls')):
                return f"[Excel file: {filename}, Size: {len(content)} bytes]"

            else:
                # Try to decode as text
                try:
                    return content.decode('utf-8', errors='ignore')
                except:
                    return f"[Binary file: {filename}, Size: {len(content)} bytes]"

        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {e}")
            return f"[Error extracting content from {filename}]"

    def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF content"""
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = pypdf.PdfReader(pdf_file)

            text_parts = []
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            return '\n'.join(text_parts)

        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            return "[PDF extraction failed]"

    def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX content"""
        try:
            doc_file = io.BytesIO(content)
            doc = Document(doc_file)

            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text_parts.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text:
                            row_text.append(cell.text)
                    if row_text:
                        text_parts.append(' | '.join(row_text))

            return '\n'.join(text_parts)

        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            return "[DOCX extraction failed]"

    async def _extract_metadata(self, documents: Dict[str, str], categorized: Dict[str, List], notice_id: str) -> Dict[str, Any]:
        """Extract structured metadata from documents"""
        metadata = {
            'notice_id': notice_id,
            'has_info_doc': len(categorized.get('info_doc', [])) > 0,
            'has_main_rfp': len(categorized.get('main_rfp', [])) > 0,
            'has_technical_specs': len(categorized.get('technical_specs', [])) > 0,
            'document_categories': {k: len(v) for k, v in categorized.items() if v}
        }

        # Try to extract structured data from info doc
        info_docs = categorized.get('info_doc', [])
        if info_docs:
            for info_doc in info_docs:
                doc_key = f"info_doc_{info_doc['name']}"
                if doc_key in documents:
                    content = documents[doc_key]

                    # Extract key metadata fields
                    metadata.update(self._parse_info_document(content))
                    break

        # Calculate content hash for deduplication
        all_content = ' '.join(str(doc) for doc in documents.values())
        metadata['content_hash'] = hashlib.sha256(all_content.encode()).hexdigest()[:16]

        return metadata

    def _parse_info_document(self, content: str) -> Dict[str, Any]:
        """Parse structured information from info document"""
        extracted = {}

        # Common patterns in info documents
        patterns = {
            'title': r'(?:Title|Name|Subject):\s*(.+?)(?:\n|$)',
            'solicitation_number': r'(?:Solicitation Number|Sol Number|Number):\s*(.+?)(?:\n|$)',
            'agency': r'(?:Agency|Organization|Department):\s*(.+?)(?:\n|$)',
            'due_date': r'(?:Due Date|Deadline|Response Date):\s*(.+?)(?:\n|$)',
            'type': r'(?:Type|Notice Type|Opportunity Type):\s*(.+?)(?:\n|$)',
            'naics_code': r'(?:NAICS|NAICS Code):\s*(\d+)',
            'set_aside': r'(?:Set[- ]Aside|Small Business):\s*(.+?)(?:\n|$)',
            'place_of_performance': r'(?:Place of Performance|Location):\s*(.+?)(?:\n|$)',
        }

        for field, pattern in patterns.items():
            match = re.search(pattern, content, re.IGNORECASE)
            if match:
                extracted[field] = match.group(1).strip()

        # Try to parse as JSON if it looks like JSON
        if content.strip().startswith('{'):
            try:
                json_data = json.loads(content)
                extracted.update(json_data)
            except:
                pass

        return extracted

    def cleanup(self):
        """Clean up temporary files"""
        try:
            import shutil
            if self.temp_dir.exists():
                shutil.rmtree(self.temp_dir)
                logger.info(f"Cleaned up temp directory: {self.temp_dir}")
        except Exception as e:
            logger.error(f"Error cleaning up temp files: {e}")