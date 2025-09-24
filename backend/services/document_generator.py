import os
import logging
from typing import Dict, Any, Optional
from pathlib import Path
import tempfile
from datetime import datetime
import markdown
import json
from jinja2 import Template
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill

logger = logging.getLogger(__name__)

class DocumentGenerator:
    def __init__(self):
        self.output_dir = Path(tempfile.mkdtemp(prefix="rfp_output_"))
        
    async def generate_document(self, content: str, doc_type: str, metadata: Dict[str, Any] = None) -> Path:
        """Generate a document in the specified format"""
        try:
            doc_name = metadata.get('name', f"document_{datetime.now().strftime('%Y%m%d_%H%M%S')}")
            
            if doc_type == 'docx' or doc_type == 'doc':
                return await self.generate_docx(content, doc_name, metadata)
            elif doc_type == 'xlsx' or doc_type == 'xls':
                return await self.generate_xlsx(content, doc_name, metadata)
            elif doc_type == 'pdf':
                return await self.generate_pdf(content, doc_name, metadata)
            elif doc_type == 'md':
                return await self.generate_markdown(content, doc_name, metadata)
            else:
                return await self.generate_text(content, doc_name, metadata)
                
        except Exception as e:
            logger.error(f"Error generating document: {e}")
            raise
    
    async def generate_docx(self, content: str, doc_name: str, metadata: Dict[str, Any] = None) -> Path:
        """Generate a DOCX document"""
        try:
            doc = Document()
            
            # Add metadata
            doc.core_properties.title = metadata.get('title', doc_name)
            doc.core_properties.author = "Kamiwaza"
            doc.core_properties.subject = metadata.get('subject', 'RFP Response')
            
            # Add header
            header = doc.sections[0].header
            header_para = header.paragraphs[0]
            header_para.text = f"Kamiwaza - {metadata.get('rfp_id', 'RFP Response')}"
            header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Parse content and add to document
            lines = content.split('\n')
            for line in lines:
                if line.startswith('# '):
                    # Main heading
                    heading = doc.add_heading(line[2:], level=1)
                    self._style_heading(heading)
                elif line.startswith('## '):
                    # Sub heading
                    heading = doc.add_heading(line[3:], level=2)
                    self._style_heading(heading)
                elif line.startswith('### '):
                    # Sub-sub heading
                    heading = doc.add_heading(line[4:], level=3)
                    self._style_heading(heading)
                elif line.startswith('- ') or line.startswith('* '):
                    # Bullet point
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line.startswith('1. ') or line[0:2] in ['2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.']:
                    # Numbered list
                    doc.add_paragraph(line[3:], style='List Number')
                elif line.strip():
                    # Regular paragraph
                    para = doc.add_paragraph(line)
                    para.paragraph_format.space_after = Pt(6)
                else:
                    # Empty line
                    doc.add_paragraph()
            
            # Add footer with page numbers
            footer = doc.sections[0].footer
            footer_para = footer.paragraphs[0]
            footer_para.text = "Page "
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Save document
            output_path = self.output_dir / f"{doc_name}.docx"
            doc.save(output_path)
            
            logger.info(f"Generated DOCX: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating DOCX: {e}")
            raise
    
    def _style_heading(self, heading):
        """Apply styling to heading"""
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    
    async def generate_xlsx(self, content: str, doc_name: str, metadata: Dict[str, Any] = None) -> Path:
        """Generate an Excel document"""
        try:
            wb = openpyxl.Workbook()
            
            # Try to parse content as structured data
            try:
                data = json.loads(content)
                if isinstance(data, dict) and 'sheets' in data:
                    # Multiple sheets
                    for sheet_name, sheet_data in data['sheets'].items():
                        ws = wb.create_sheet(title=sheet_name[:31])
                        await self._populate_sheet(ws, sheet_data)
                    # Remove default sheet if we created others
                    if len(wb.worksheets) > 1:
                        wb.remove(wb.worksheets[0])
                else:
                    # Single sheet with data
                    ws = wb.active
                    ws.title = metadata.get('sheet_name', 'Data')[:31]
                    await self._populate_sheet(ws, data)
            except json.JSONDecodeError:
                # Fallback: treat as text content
                ws = wb.active
                ws.title = "Content"
                lines = content.split('\n')
                for row_idx, line in enumerate(lines, 1):
                    if '\t' in line:
                        # Tab-separated values
                        values = line.split('\t')
                        for col_idx, value in enumerate(values, 1):
                            ws.cell(row=row_idx, column=col_idx, value=value)
                    else:
                        ws.cell(row=row_idx, column=1, value=line)
            
            # Style the first row as header
            for cell in wb.active[1]:
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
                cell.alignment = Alignment(horizontal="center")
            
            # Auto-adjust column widths
            for column in wb.active.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                wb.active.column_dimensions[column_letter].width = adjusted_width
            
            # Save workbook
            output_path = self.output_dir / f"{doc_name}.xlsx"
            wb.save(output_path)
            
            logger.info(f"Generated XLSX: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating XLSX: {e}")
            raise
    
    async def _populate_sheet(self, worksheet, data):
        """Populate worksheet with data"""
        if isinstance(data, list):
            # List of records
            if data and isinstance(data[0], dict):
                # List of dictionaries - create table
                headers = list(data[0].keys())
                for col_idx, header in enumerate(headers, 1):
                    worksheet.cell(row=1, column=col_idx, value=header)
                
                for row_idx, record in enumerate(data, 2):
                    for col_idx, header in enumerate(headers, 1):
                        worksheet.cell(row=row_idx, column=col_idx, value=record.get(header, ''))
            else:
                # Simple list
                for row_idx, item in enumerate(data, 1):
                    worksheet.cell(row=row_idx, column=1, value=str(item))
        
        elif isinstance(data, dict):
            # Dictionary - key-value pairs
            row_idx = 1
            for key, value in data.items():
                worksheet.cell(row=row_idx, column=1, value=key)
                worksheet.cell(row=row_idx, column=2, value=str(value))
                row_idx += 1
    
    async def generate_pdf(self, content: str, doc_name: str, metadata: Dict[str, Any] = None) -> Path:
        """Generate a PDF document"""
        try:
            # First generate as markdown/HTML then convert to PDF
            html_content = markdown.markdown(content, extensions=['tables', 'fenced_code'])
            
            # Add CSS styling
            styled_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="utf-8">
                <title>{metadata.get('title', doc_name)}</title>
                <style>
                    body {{ 
                        font-family: 'Helvetica', 'Arial', sans-serif; 
                        line-height: 1.6;
                        color: #333;
                        max-width: 800px;
                        margin: 0 auto;
                        padding: 20px;
                    }}
                    h1 {{ color: #003366; border-bottom: 2px solid #003366; padding-bottom: 10px; }}
                    h2 {{ color: #004080; margin-top: 30px; }}
                    h3 {{ color: #0066cc; }}
                    table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
                    th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                    th {{ background-color: #f2f2f2; font-weight: bold; }}
                    .header {{ 
                        text-align: center; 
                        margin-bottom: 40px;
                        padding: 20px;
                        background-color: #f8f9fa;
                    }}
                    .footer {{ 
                        text-align: center; 
                        margin-top: 40px;
                        padding-top: 20px;
                        border-top: 1px solid #ddd;
                        color: #666;
                    }}
                </style>
            </head>
            <body>
                <div class="header">
                    <h1>Kamiwaza</h1>
                    <p>{metadata.get('subtitle', 'RFP Response Document')}</p>
                </div>
                {html_content}
                <div class="footer">
                    <p>© {datetime.now().year} Kamiwaza. Confidential and Proprietary.</p>
                </div>
            </body>
            </html>
            """
            
            # For now, save as HTML (PDF conversion would require additional libraries like weasyprint)
            output_path = self.output_dir / f"{doc_name}.html"
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(styled_html)
            
            logger.info(f"Generated PDF (HTML): {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating PDF: {e}")
            raise
    
    async def generate_markdown(self, content: str, doc_name: str, metadata: Dict[str, Any] = None) -> Path:
        """Generate a Markdown document"""
        try:
            # Add metadata header
            header = f"""---
title: {metadata.get('title', doc_name)}
author: Kamiwaza
date: {datetime.now().strftime('%Y-%m-%d')}
subject: {metadata.get('subject', 'RFP Response')}
---

# {metadata.get('title', doc_name)}

"""
            
            # Combine header with content
            full_content = header + content
            
            # Save markdown file
            output_path = self.output_dir / f"{doc_name}.md"
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(full_content)
            
            logger.info(f"Generated Markdown: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating Markdown: {e}")
            raise
    
    async def generate_text(self, content: str, doc_name: str, metadata: Dict[str, Any] = None) -> Path:
        """Generate a plain text document"""
        try:
            # Add simple header
            header = f"""{'=' * 60}
{metadata.get('title', doc_name).upper()}
Kamiwaza RFP Response
Date: {datetime.now().strftime('%Y-%m-%d')}
{'=' * 60}

"""
            
            # Combine header with content
            full_content = header + content
            
            # Save text file
            output_path = self.output_dir / f"{doc_name}.txt"
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(full_content)
            
            logger.info(f"Generated Text: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating text document: {e}")
            raise
    
    def cleanup(self):
        """Clean up temporary output directory"""
        try:
            import shutil
            if self.output_dir.exists():
                shutil.rmtree(self.output_dir)
        except Exception as e:
            logger.warning(f"Error cleaning up output directory: {e}")